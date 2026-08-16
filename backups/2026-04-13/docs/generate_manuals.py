#!/usr/bin/env python3
"""Generate FairwayConnect Member and Admin user manuals as .docx files."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── Constants ──
DARK_GREEN = RGBColor(0x1B, 0x5E, 0x3A)
LIGHT_GREEN = RGBColor(0xE8, 0xF5, 0xE9)
MID_GREEN = RGBColor(0x2E, 0x7D, 0x32)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
GREY = RGBColor(0x66, 0x66, 0x66)
FONT_NAME = 'Calibri'
OUTPUT_DIR = '/Users/abcooney/.openclaw/workspace/fairway-connect/docs'


# ── Helpers ──

def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    """Set borders on a cell. kwargs: top, bottom, left, right with dict of sz, color, val."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, attrs in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{attrs.get("val", "single")}" '
            f'w:sz="{attrs.get("sz", "4")}" w:space="0" w:color="{attrs.get("color", "000000")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def add_table_borders(table):
    """Add borders to all cells in a table."""
    border_attrs = {"val": "single", "sz": "4", "color": "CCCCCC"}
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top=border_attrs, bottom=border_attrs,
                            left=border_attrs, right=border_attrs)


def styled_paragraph(doc, text, font_size=11, bold=False, color=BLACK, alignment=None, space_after=Pt(6), space_before=Pt(0), font_name=FONT_NAME):
    """Add a paragraph with specific styling."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font_name
    if alignment:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = space_after
    pf.space_before = space_before
    return p


def add_heading_styled(doc, text, level=1):
    """Add a heading with dark green color."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = DARK_GREEN
        run.font.name = FONT_NAME
    return h


def add_callout_box(doc, title, text, box_type="note"):
    """Add a callout box (Note/Important/Tip)."""
    icons = {"note": "📝", "important": "⚠️", "tip": "💡"}
    colors = {"note": "E3F2FD", "important": "FFF3E0", "tip": "E8F5E9"}
    icon = icons.get(box_type, "📝")
    bg_color = colors.get(box_type, "E3F2FD")

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, bg_color)

    # Title
    p = cell.paragraphs[0]
    run = p.add_run(f"{icon} {title}")
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.name = FONT_NAME
    run.font.color.rgb = DARK_GREEN

    # Body
    p2 = cell.add_paragraph()
    run2 = p2.add_run(text)
    run2.font.size = Pt(10)
    run2.font.name = FONT_NAME
    run2.font.color.rgb = BLACK

    # Set width
    for row in table.rows:
        for c in row.cells:
            c.width = Cm(16)

    border_attrs = {"val": "single", "sz": "6", "color": "1B5E3A"}
    set_cell_border(cell, left=border_attrs)

    doc.add_paragraph()  # spacer


def add_what_youll_see(doc, items):
    """Add a 'What you'll see' box for the member guide."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "E8F5E9")

    p = cell.paragraphs[0]
    run = p.add_run("👀 What you'll see:")
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.name = FONT_NAME
    run.font.color.rgb = DARK_GREEN

    for item in items:
        p2 = cell.add_paragraph()
        p2.style = 'List Bullet'
        run2 = p2.add_run(item)
        run2.font.size = Pt(10)
        run2.font.name = FONT_NAME

    for row in table.rows:
        for c in row.cells:
            c.width = Cm(16)

    border_attrs = {"val": "single", "sz": "6", "color": "2E7D32"}
    set_cell_border(cell, left=border_attrs)
    doc.add_paragraph()


def setup_document(doc, footer_text):
    """Setup page size, margins, default font, headers, footers."""
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(11)
    font.color.rgb = BLACK

    # Heading styles
    for i in range(1, 4):
        hs = doc.styles[f'Heading {i}']
        hs.font.name = FONT_NAME
        hs.font.color.rgb = DARK_GREEN

    # Page setup - A4
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # Footer
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add footer text
    run = fp.add_run(footer_text + "    |    Page ")
    run.font.size = Pt(8)
    run.font.color.rgb = GREY
    run.font.name = FONT_NAME

    # Add page number field
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    fp.runs[-1]._r.addnext(fldChar1)
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    fldChar1.addnext(instrText)
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    instrText.addnext(fldChar2)


def add_title_page(doc, title, subtitle, version):
    """Add a branded title page with green theme."""
    # Add several blank paragraphs for spacing
    for _ in range(6):
        doc.add_paragraph()

    # Green bar (simulated with a table)
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "1B5E3A")
    cell.width = Cm(16)

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # App name
    run = p.add_run("⛳ FairwayConnect")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = WHITE
    run.font.name = FONT_NAME

    # Subtitle
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(title)
    run2.font.size = Pt(24)
    run2.font.color.rgb = WHITE
    run2.font.name = FONT_NAME

    # Version
    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(version)
    run3.font.size = Pt(14)
    run3.font.color.rgb = RGBColor(0xA5, 0xD6, 0xA7)
    run3.font.name = FONT_NAME

    # Extra padding paragraph in cell
    p4 = cell.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run(subtitle)
    run4.font.size = Pt(12)
    run4.font.color.rgb = RGBColor(0xC8, 0xE6, 0xC9)
    run4.font.name = FONT_NAME

    # Spacing after
    for _ in range(4):
        doc.add_paragraph()

    # Date
    styled_paragraph(doc, "April 2026", font_size=12, color=GREY,
                     alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Page break
    doc.add_page_break()


def add_toc(doc):
    """Add a Table of Contents field."""
    add_heading_styled(doc, "Table of Contents", level=1)

    p = doc.add_paragraph()
    run = p.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar1)

    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
    run._r.append(instrText)

    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run._r.append(fldChar2)

    # Placeholder text
    run2 = p.add_run("Right-click and select 'Update Field' to generate table of contents")
    run2.font.size = Pt(10)
    run2.font.color.rgb = GREY
    run2.font.name = FONT_NAME
    run2.font.italic = True

    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run2._r.append(fldChar3)

    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════
# DOCUMENT 1: MEMBER GUIDE
# ═══════════════════════════════════════════════════════════════

def create_member_guide():
    doc = Document()
    setup_document(doc, "FairwayConnect - Member Guide v1.0")
    add_title_page(doc, "Member Guide", "Your complete guide to using FairwayConnect", "Version 1.0")
    add_toc(doc)

    # ── 1. Welcome ──
    add_heading_styled(doc, "1. Welcome to FairwayConnect ⛳", level=1)
    styled_paragraph(doc,
        "Welcome to FairwayConnect — your golf society's digital home! 🏌️‍♂️",
        font_size=12, bold=True, color=DARK_GREEN)
    styled_paragraph(doc,
        "FairwayConnect is a web app built specifically for your golf society. "
        "It's where you'll find everything about your outings — from upcoming events "
        "and live leaderboards to results, prizes, and the all-important Golfer of the Year standings.")
    styled_paragraph(doc, "")

    styled_paragraph(doc, "As a member, you can:", font_size=11, bold=True)
    bullets = [
        "📅  Check the full season calendar and see what's coming up",
        "📊  Follow live scores during an event",
        "🏆  See who's leading the Golfer of the Year race",
        "📋  View results, prizes, and side competition winners",
        "🔢  Check the deductions sheet to see where you stand",
    ]
    for b in bullets:
        p = doc.add_paragraph(b, style='List Bullet')
        for run in p.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(11)

    styled_paragraph(doc, "")
    styled_paragraph(doc,
        "No accounts, no passwords, no fuss. Just open the link and you're in! 👍",
        font_size=11, bold=True, color=MID_GREEN)

    doc.add_page_break()

    # ── 2. Getting Started ──
    add_heading_styled(doc, "2. Getting Started 🚀", level=1)
    styled_paragraph(doc,
        "Getting started with FairwayConnect couldn't be easier. There's no app to download "
        "and no account to create.")
    styled_paragraph(doc, "")

    add_heading_styled(doc, "How to Access", level=2)

    # Step table
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(table)

    steps = [
        ("Step", "What to Do"),
        ("1️⃣", "Open any web browser on your phone, tablet, or computer"),
        ("2️⃣", "Go to: fairwayconnect.fly.dev"),
        ("3️⃣", "That's it — you're in! No login required 🎉"),
    ]
    for i, (step, action) in enumerate(steps):
        row = table.rows[i]
        row.cells[0].width = Cm(2)
        row.cells[1].width = Cm(14)

        p0 = row.cells[0].paragraphs[0]
        r0 = p0.add_run(step)
        r0.font.name = FONT_NAME
        r0.font.size = Pt(10)
        r0.font.bold = (i == 0)
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p1 = row.cells[1].paragraphs[0]
        r1 = p1.add_run(action)
        r1.font.name = FONT_NAME
        r1.font.size = Pt(10)
        r1.font.bold = (i == 0)

        if i == 0:
            set_cell_shading(row.cells[0], "1B5E3A")
            set_cell_shading(row.cells[1], "1B5E3A")
            r0.font.color.rgb = WHITE
            r1.font.color.rgb = WHITE

    styled_paragraph(doc, "")
    add_callout_box(doc, "Tip", "Add the site to your phone's home screen for quick access! "
                    "On iPhone: tap Share → Add to Home Screen. On Android: tap the menu → Add to Home Screen.",
                    box_type="tip")

    doc.add_page_break()

    # ── 3. Home Page ──
    add_heading_styled(doc, "3. Home Page 🏠", level=1)
    styled_paragraph(doc,
        "The home page is your dashboard — a quick snapshot of everything that matters right now.")

    add_what_youll_see(doc, [
        "🗓️  Next Event — The upcoming outing with date, course name, and details",
        "🏆  OOM Leader — The current leader in the Golfer of the Year (Order of Merit) race",
        "📊  Recent Results — A summary of the most recent outing results",
        "⚡  Quick navigation to Calendar, Leaderboard, and GOTY pages",
    ])

    styled_paragraph(doc,
        "The home page updates automatically, so every time you visit you'll see the latest information. "
        "It's designed to give you a quick glance at what's happening in the society without having to dig around.")

    doc.add_page_break()

    # ── 4. Calendar ──
    add_heading_styled(doc, "4. Calendar 📅", level=1)
    styled_paragraph(doc,
        "The Calendar page shows every event in the current season, laid out in a clear, "
        "easy-to-read format.")

    add_what_youll_see(doc, [
        "📋  A list of all outings for the season (past and upcoming)",
        "📍  Course name and date for each event",
        "✅  Status badges — see which events are completed, in progress, or upcoming",
        "🔗  Click any event to view its details, scorecard, and results",
    ])

    add_heading_styled(doc, "How to Use It", level=2)
    styled_paragraph(doc,
        "Simply scroll through the list to see all events. Past events will show their results, "
        "while upcoming events will show the date and venue.")
    styled_paragraph(doc, "")
    styled_paragraph(doc,
        "Tap or click on any event to dive into the details — you'll see the full scorecard, "
        "prizes, and side competition results for completed events.",
        font_size=11)

    add_callout_box(doc, "Tip", "Keep an eye on the calendar to plan ahead! "
                    "You can see the full season at a glance and never miss an outing.",
                    box_type="tip")

    doc.add_page_break()

    # ── 5. Leaderboard ──
    add_heading_styled(doc, "5. Leaderboard 📊", level=1)
    styled_paragraph(doc,
        "The Leaderboard is where the magic happens during an event! 🔥 "
        "As scores are entered, you can follow along in real time.")

    add_what_youll_see(doc, [
        "🥇  Player rankings based on Stableford points",
        "📈  Live scores updating as they're entered by the admin",
        "🏌️  Front 9 and Back 9 standings — separate leaderboards for each half",
        "📊  Overall standings with running totals",
        "🎯  Points for each hole (when scores have been entered)",
    ])

    add_heading_styled(doc, "During an Event", level=2)
    styled_paragraph(doc,
        "While an outing is in progress, the leaderboard updates as the admin enters scores. "
        "You can check in throughout the day to see how everyone is doing.")
    styled_paragraph(doc, "")
    styled_paragraph(doc,
        "The leaderboard shows three views:", font_size=11, bold=True)
    bullets = [
        "Overall — Total Stableford points for all 18 holes",
        "Front 9 — Points for holes 1-9 only",
        "Back 9 — Points for holes 10-18 only",
    ]
    for b in bullets:
        p = doc.add_paragraph(b, style='List Bullet')
        for run in p.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(11)

    add_callout_box(doc, "Note", "Scores are entered by the event admin, so there may be a slight "
                    "delay between a player finishing a hole and the score appearing on the leaderboard.",
                    box_type="note")

    doc.add_page_break()

    # ── 6. GOTY ──
    add_heading_styled(doc, "6. Golfer of the Year (GOTY) 🏆", level=1)
    styled_paragraph(doc,
        "The Golfer of the Year — also known as the Order of Merit (OOM) — is the big one! "
        "This is the season-long competition that crowns the society's best golfer.")

    add_what_youll_see(doc, [
        "🏆  Current GOTY standings with all players ranked",
        "📊  Points from each outing shown in columns",
        "🔢  Running totals updated after each event",
        "⭐  Best 6 of 8 scores highlighted",
    ])

    add_heading_styled(doc, "How It Works", level=2)

    # Rules table
    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(table)

    rules = [
        ("Rule", "Detail"),
        ("Scoring", "Your Stableford points from each outing count towards GOTY"),
        ("Best Scores", "Your best 6 scores out of 8 outings count"),
        ("Ranking", "Players are ranked by total points from their best 6 rounds"),
        ("No Deductions", "GOTY uses raw Stableford points — deductions don't apply here"),
    ]
    for i, (rule, detail) in enumerate(rules):
        row = table.rows[i]
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(12)

        p0 = row.cells[0].paragraphs[0]
        r0 = p0.add_run(rule)
        r0.font.name = FONT_NAME
        r0.font.size = Pt(10)
        r0.font.bold = True

        p1 = row.cells[1].paragraphs[0]
        r1 = p1.add_run(detail)
        r1.font.name = FONT_NAME
        r1.font.size = Pt(10)

        if i == 0:
            set_cell_shading(row.cells[0], "1B5E3A")
            set_cell_shading(row.cells[1], "1B5E3A")
            r0.font.color.rgb = WHITE
            r1.font.color.rgb = WHITE
        elif i % 2 == 0:
            set_cell_shading(row.cells[0], "E8F5E9")
            set_cell_shading(row.cells[1], "E8F5E9")

    styled_paragraph(doc, "")
    styled_paragraph(doc,
        "The GOTY standings update automatically after each event is finalised. "
        "Check back regularly to see where you stand in the race! 📈")

    doc.add_page_break()

    # ── 7. Results ──
    add_heading_styled(doc, "7. Results 📋", level=1)
    styled_paragraph(doc,
        "After each outing is completed and published, you can view the full results "
        "including prizes, side competitions, and complete scorecards.")

    add_what_youll_see(doc, [
        "🏅  Prize winners — 1st, 2nd, 3rd place",
        "⭐  Front 9 and Back 9 winners",
        "🎯  Side competition results (Nearest the Pin, Longest Drive, Twos Club)",
        "📊  Full scorecard with every player's hole-by-hole scores and Stableford points",
        "📉  Deductions applied to each player's score",
    ])

    add_heading_styled(doc, "Viewing Results", level=2)
    styled_paragraph(doc,
        "To view results for any outing:")
    steps = [
        "Go to the Calendar page",
        "Click on a completed event (marked with a ✅ status)",
        "You'll see the results page with tabs for Scores, Prizes, and Side Competitions",
    ]
    for i, s in enumerate(steps, 1):
        p = doc.add_paragraph(f"{i}. {s}")
        for run in p.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(11)

    styled_paragraph(doc, "")
    add_heading_styled(doc, "Understanding Prizes", level=2)
    styled_paragraph(doc,
        "Prizes are awarded based on your Stableford score after deductions are applied. "
        "The Prizes tab shows who won and how deductions affected the final standings.")

    doc.add_page_break()

    # ── 8. Deductions ──
    add_heading_styled(doc, "8. Deductions 🔢", level=1)
    styled_paragraph(doc,
        "Deductions are a key part of how the society keeps things competitive and fair. "
        "If you've been winning, you'll have shots deducted from your score in future events!")

    add_heading_styled(doc, "What Are Deductions?", level=2)
    styled_paragraph(doc,
        "Deductions are a society rule designed to level the playing field. When you win a prize, "
        "you get shots deducted from your Stableford score in the next event. "
        "Think of it as a penalty for being too good! 😄")

    add_heading_styled(doc, "How They Work", level=2)

    # Deductions table
    table = doc.add_table(rows=6, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(table)

    deductions = [
        ("Prize Won", "Deduction Applied"),
        ("🥇 1st Place", "-3 shots from next event"),
        ("🥈 2nd Place", "-2 shots from next event"),
        ("🥉 3rd Place", "-2 shots from next event"),
        ("⭐ Front 9 / Back 9", "-1 shot from next event"),
        ("Played, no prize", "+1 shot earned back"),
    ]
    for i, (prize, ded) in enumerate(deductions):
        row = table.rows[i]
        row.cells[0].width = Cm(6)
        row.cells[1].width = Cm(10)

        p0 = row.cells[0].paragraphs[0]
        r0 = p0.add_run(prize)
        r0.font.name = FONT_NAME
        r0.font.size = Pt(10)
        r0.font.bold = (i == 0)

        p1 = row.cells[1].paragraphs[0]
        r1 = p1.add_run(ded)
        r1.font.name = FONT_NAME
        r1.font.size = Pt(10)
        r1.font.bold = (i == 0)

        if i == 0:
            set_cell_shading(row.cells[0], "1B5E3A")
            set_cell_shading(row.cells[1], "1B5E3A")
            r0.font.color.rgb = WHITE
            r1.font.color.rgb = WHITE
        elif i % 2 == 0:
            set_cell_shading(row.cells[0], "E8F5E9")
            set_cell_shading(row.cells[1], "E8F5E9")

    styled_paragraph(doc, "")
    add_heading_styled(doc, "Reading the Deductions Sheet", level=2)
    styled_paragraph(doc,
        "The deductions sheet shows each player's running total of deductions. "
        "You'll see:")
    bullets = [
        "Your starting deductions for the year",
        "Changes from each outing (wins add deductions, playing earns one back)",
        "Your current deduction total",
    ]
    for b in bullets:
        p = doc.add_paragraph(b, style='List Bullet')
        for run in p.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(11)

    add_callout_box(doc, "Important", "Deductions only affect the overall prize rankings. "
                    "They do NOT apply to Front 9/Back 9 prizes or GOTY standings. "
                    "So your Golfer of the Year score is always your pure Stableford total!",
                    box_type="important")

    styled_paragraph(doc, "")
    styled_paragraph(doc, "")

    # ── End ──
    doc.add_page_break()
    add_heading_styled(doc, "Need Help? 🤝", level=1)
    styled_paragraph(doc,
        "FairwayConnect is designed to be simple and intuitive. If you have any questions "
        "or spot something that doesn't look right, just let your society admin know.")
    styled_paragraph(doc, "")
    styled_paragraph(doc,
        "Now get out there and make some birdies! 🏌️‍♂️⛳",
        font_size=14, bold=True, color=DARK_GREEN, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Save
    path = os.path.join(OUTPUT_DIR, 'FairwayConnect-Member-Guide-v1.0.docx')
    doc.save(path)
    print(f"✅ Member Guide saved to: {path}")
    return path


# ═══════════════════════════════════════════════════════════════
# DOCUMENT 2: ADMIN GUIDE
# ═══════════════════════════════════════════════════════════════

def create_admin_guide():
    doc = Document()
    setup_document(doc, "FairwayConnect - Admin Guide v1.0")
    add_title_page(doc, "Admin Guide", "Complete administration guide for managing your golf society", "Version 1.0")
    add_toc(doc)

    # ── 1. Getting Started ──
    add_heading_styled(doc, "1. Getting Started", level=1)
    styled_paragraph(doc,
        "This guide covers everything you need to manage your golf society through the "
        "FairwayConnect admin dashboard. From creating events to publishing results, "
        "this manual will walk you through each step.")

    add_heading_styled(doc, "Accessing the Admin Dashboard", level=2)
    styled_paragraph(doc, "To access the admin area:")

    steps = [
        "Open your browser and go to fairwayconnect.fly.dev/admin",
        "Enter the admin PIN: 2026",
        "You'll be taken to the admin dashboard",
    ]
    for i, s in enumerate(steps, 1):
        p = doc.add_paragraph(f"{i}. {s}")
        for run in p.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(11)

    add_callout_box(doc, "Important", "Keep the admin PIN secure. Only share it with authorised "
                    "administrators. The PIN provides full access to create events, enter scores, "
                    "and publish results.", box_type="important")

    add_heading_styled(doc, "Admin Dashboard Overview", level=2)
    styled_paragraph(doc,
        "The admin dashboard is your central hub. From here you can:")

    features = [
        ("Create & manage events", "Set up outings with course details, dates, and tee times"),
        ("Manage players", "Add players to events, track RSVPs"),
        ("Enter scores", "Input hole-by-hole scores during or after an outing"),
        ("Publish results", "Finalise events, generate prizes, and update standings"),
        ("Manage deductions", "View and edit the deductions sheet"),
        ("Update handicaps", "Keep member handicap indexes current"),
    ]
    table = doc.add_table(rows=len(features) + 1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(table)

    # Header
    for j, header in enumerate(["Feature", "Description"]):
        cell = table.rows[0].cells[j]
        set_cell_shading(cell, "1B5E3A")
        p = cell.paragraphs[0]
        r = p.add_run(header)
        r.font.bold = True
        r.font.color.rgb = WHITE
        r.font.name = FONT_NAME
        r.font.size = Pt(10)

    for i, (feat, desc) in enumerate(features):
        row = table.rows[i + 1]
        p0 = row.cells[0].paragraphs[0]
        r0 = p0.add_run(feat)
        r0.font.name = FONT_NAME
        r0.font.size = Pt(10)
        r0.font.bold = True

        p1 = row.cells[1].paragraphs[0]
        r1 = p1.add_run(desc)
        r1.font.name = FONT_NAME
        r1.font.size = Pt(10)

        if i % 2 == 0:
            set_cell_shading(row.cells[0], "E8F5E9")
            set_cell_shading(row.cells[1], "E8F5E9")

    doc.add_page_break()

    # ── 2. Managing Events ──
    add_heading_styled(doc, "2. Managing Events", level=1)
    styled_paragraph(doc,
        "Events are the core of FairwayConnect. Each outing is set up as an event with "
        "course details, date, and tee time configuration.")

    add_heading_styled(doc, "2.1 Creating a New Event", level=2)
    styled_paragraph(doc, "To create a new event:")
    steps = [
        "From the admin dashboard, click 'Create Event' or the '+' button",
        "Enter the event details:\n    • Event name (e.g., 'Spring Outing - Portmarnock')\n    • Date of the outing\n    • Course name",
        "Configure the tee settings (see Section 2.2)",
        "Save the event — it will appear on the calendar immediately",
    ]
    for i, s in enumerate(steps, 1):
        p = doc.add_paragraph(f"{i}. {s}")
        for run in p.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(11)

    add_heading_styled(doc, "2.2 WHS Course Settings", level=2)
    styled_paragraph(doc,
        "For accurate handicap calculations, you need to set up the course details "
        "according to the World Handicap System (WHS).")

    table = doc.add_table(rows=6, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(table)

    settings = [
        ("Setting", "Description"),
        ("Slope Rating", "The slope rating of the course from the selected tees (e.g., 128)"),
        ("Course Rating (CR)", "The course rating for the selected tees (e.g., 72.1)"),
        ("Par", "The par for the course (e.g., 72)"),
        ("Tees", "Select which tees will be played (e.g., White, Yellow, Red)"),
    ]
    for i, (setting, desc) in enumerate(settings):
        row = table.rows[i]
        p0 = row.cells[0].paragraphs[0]
        r0 = p0.add_run(setting)
        r0.font.name = FONT_NAME
        r0.font.size = Pt(10)
        r0.font.bold = True

        p1 = row.cells[1].paragraphs[0]
        r1 = p1.add_run(desc)
        r1.font.name = FONT_NAME
        r1.font.size = Pt(10)

        if i == 0:
            set_cell_shading(row.cells[0], "1B5E3A")
            set_cell_shading(row.cells[1], "1B5E3A")
            r0.font.color.rgb = WHITE
            r1.font.color.rgb = WHITE
        elif i % 2 == 0:
            set_cell_shading(row.cells[0], "E8F5E9")
            set_cell_shading(row.cells[1], "E8F5E9")

    styled_paragraph(doc, "")

    add_heading_styled(doc, "2.3 Handicap Allowance", level=2)
    styled_paragraph(doc,
        "FairwayConnect uses a 95% handicap allowance as standard, in line with Golf Ireland "
        "recommendations for individual Stableford competitions.")

    add_callout_box(doc, "Note", "The 95% allowance is applied automatically when calculating "
                    "playing handicaps. You don't need to adjust this manually — the system "
                    "handles it based on each player's handicap index and the course settings.",
                    box_type="note")

    doc.add_page_break()

    # ── 3. Players & RSVPs ──
    add_heading_styled(doc, "3. Players & RSVPs", level=1)

    add_heading_styled(doc, "3.1 Adding Players to Events", level=2)
    styled_paragraph(doc, "To add players to an event:")
    steps = [
        "Open the event from the admin dashboard",
        "Go to the Players tab",
        "Click 'Add Player' and select from the member list",
        "The player will be added with a default RSVP status",
    ]
    for i, s in enumerate(steps, 1):
        p = doc.add_paragraph(f"{i}. {s}")
        for run in p.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(11)

    add_heading_styled(doc, "3.2 Managing RSVPs", level=2)
    styled_paragraph(doc,
        "Each player's RSVP status can be set to one of three states:")

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(table)

    rsvps = [
        ("Status", "Meaning"),
        ("✅ Confirmed", "Player has confirmed they will attend"),
        ("❓ Maybe", "Player hasn't decided yet"),
        ("❌ Declined", "Player has declined the invitation"),
    ]
    for i, (status, meaning) in enumerate(rsvps):
        row = table.rows[i]
        p0 = row.cells[0].paragraphs[0]
        r0 = p0.add_run(status)
        r0.font.name = FONT_NAME
        r0.font.size = Pt(10)
        r0.font.bold = (i == 0)

        p1 = row.cells[1].paragraphs[0]
        r1 = p1.add_run(meaning)
        r1.font.name = FONT_NAME
        r1.font.size = Pt(10)

        if i == 0:
            set_cell_shading(row.cells[0], "1B5E3A")
            set_cell_shading(row.cells[1], "1B5E3A")
            r0.font.color.rgb = WHITE
            r1.font.color.rgb = WHITE

    styled_paragraph(doc, "")
    styled_paragraph(doc,
        "You can easily see who hasn't responded yet — players without a confirmed "
        "or declined status will be highlighted.")

    add_heading_styled(doc, "3.3 Tee Time Management", level=2)
    styled_paragraph(doc, "To set up tee times:")
    steps = [
        "Go to the Tee Times section of the event",
        "Set the first tee time and interval between groups (e.g., 10 minutes)",
        "Drag and drop players into groups",
        "Arrange the order within each group as needed",
        "Click 'Recalculate Times' if you change the interval or first tee time",
    ]
    for i, s in enumerate(steps, 1):
        p = doc.add_paragraph(f"{i}. {s}")
        for run in p.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(11)

    add_callout_box(doc, "Tip", "Use the drag-and-drop feature to quickly rearrange players "
                    "between groups. The 'Recalculate Times' button will automatically update "
                    "all tee times based on the interval you've set.",
                    box_type="tip")

    doc.add_page_break()

    # ── 4. Entering Scores ──
    add_heading_styled(doc, "4. Entering Scores", level=1)
    styled_paragraph(doc,
        "Score entry is done through the Scorecards tab on each event. Scores are entered "
        "as gross (actual shots taken) and the system automatically calculates Stableford points.")

    add_heading_styled(doc, "4.1 Using the Scorecards Tab", level=2)
    styled_paragraph(doc, "To enter scores for a player:")
    steps = [
        "Open the event and go to the Scorecards tab",
        "Select the player from the list",
        "Enter the gross score for each hole (1-18)",
        "Stableford points are calculated automatically based on the player's playing handicap",
        "The leaderboard updates in real time as scores are saved",
    ]
    for i, s in enumerate(steps, 1):
        p = doc.add_paragraph(f"{i}. {s}")
        for run in p.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(11)

    add_heading_styled(doc, "4.2 Stableford Auto-Calculation", level=2)
    styled_paragraph(doc,
        "The system automatically calculates Stableford points using the following formula:")
    styled_paragraph(doc, "")

    table = doc.add_table(rows=8, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(table)

    stableford = [
        ("Score vs Par (net)", "Stableford Points"),
        ("2+ under par (Eagle or better)", "4 points"),
        ("1 under par (Birdie)", "3 points"),
        ("Par", "2 points"),
        ("1 over par (Bogey)", "1 point"),
        ("2 over par (Double bogey)", "0 points"),
        ("3+ over par", "0 points"),
        ("No score / Pick up", "0 points"),
    ]
    for i, (score, pts) in enumerate(stableford):
        row = table.rows[i]
        p0 = row.cells[0].paragraphs[0]
        r0 = p0.add_run(score)
        r0.font.name = FONT_NAME
        r0.font.size = Pt(10)
        r0.font.bold = (i == 0)

        p1 = row.cells[1].paragraphs[0]
        r1 = p1.add_run(pts)
        r1.font.name = FONT_NAME
        r1.font.size = Pt(10)
        r1.font.bold = (i == 0)

        if i == 0:
            set_cell_shading(row.cells[0], "1B5E3A")
            set_cell_shading(row.cells[1], "1B5E3A")
            r0.font.color.rgb = WHITE
            r1.font.color.rgb = WHITE
        elif i % 2 == 0:
            set_cell_shading(row.cells[0], "E8F5E9")
            set_cell_shading(row.cells[1], "E8F5E9")

    styled_paragraph(doc, "")

    add_heading_styled(doc, "4.3 Handicap Deductions & Scoring", level=2)
    styled_paragraph(doc,
        "The playing handicap is calculated from each player's handicap index using the "
        "WHS formula with the course's slope rating, course rating, and par. The 95% "
        "allowance is applied automatically.")
    styled_paragraph(doc,
        "Handicap strokes are distributed across holes based on the stroke index. "
        "The system handles all of this automatically — you just enter the gross scores.")

    add_callout_box(doc, "Important", "Always enter GROSS scores (actual shots taken). "
                    "Do not subtract handicap strokes — the system does this automatically "
                    "when calculating Stableford points.", box_type="important")

    doc.add_page_break()

    # ── 5. Side Competitions ──
    add_heading_styled(doc, "5. Side Competitions", level=1)
    styled_paragraph(doc,
        "Side competitions add extra fun and prizes to each outing. FairwayConnect supports "
        "several types of side comps.")

    add_heading_styled(doc, "5.1 Types of Side Competitions", level=2)

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(table)

    sides = [
        ("Competition", "Description"),
        ("Nearest the Pin (NTP)", "Closest tee shot to the pin on a par 3. Measured from the pin."),
        ("Longest Drive", "Longest drive on a designated hole. Must be on the fairway."),
        ("Twos Club", "Any player who scores a 2 on any hole (gross score)."),
    ]
    for i, (comp, desc) in enumerate(sides):
        row = table.rows[i]
        p0 = row.cells[0].paragraphs[0]
        r0 = p0.add_run(comp)
        r0.font.name = FONT_NAME
        r0.font.size = Pt(10)
        r0.font.bold = True

        p1 = row.cells[1].paragraphs[0]
        r1 = p1.add_run(desc)
        r1.font.name = FONT_NAME
        r1.font.size = Pt(10)

        if i == 0:
            set_cell_shading(row.cells[0], "1B5E3A")
            set_cell_shading(row.cells[1], "1B5E3A")
            r0.font.color.rgb = WHITE
            r1.font.color.rgb = WHITE
        elif i % 2 == 0:
            set_cell_shading(row.cells[0], "E8F5E9")
            set_cell_shading(row.cells[1], "E8F5E9")

    styled_paragraph(doc, "")

    add_heading_styled(doc, "5.2 Adding Side Competition Entries", level=2)
    styled_paragraph(doc, "To add a side competition result:")
    steps = [
        "Open the event and go to the Side Competitions section",
        "Select the competition type (NTP, Longest Drive, or Twos)",
        "Select the player and the relevant hole number",
        "Add any additional details (e.g., distance for NTP)",
        "Save the entry",
    ]
    for i, s in enumerate(steps, 1):
        p = doc.add_paragraph(f"{i}. {s}")
        for run in p.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(11)

    add_heading_styled(doc, "5.3 Editing or Deleting Entries", level=2)
    styled_paragraph(doc,
        "You can edit or delete any side competition entry at any time before the event "
        "is finalised. Simply click on the entry to edit, or use the delete button to remove it.")

    doc.add_page_break()

    # ── 6. Publishing Results ──
    add_heading_styled(doc, "6. Publishing Results", level=1)
    styled_paragraph(doc,
        "Publishing results is the final step after all scores have been entered. "
        "This process generates prizes, applies deductions, and updates the GOTY standings.")

    add_heading_styled(doc, "6.1 Recalculate Scores", level=2)
    styled_paragraph(doc, "Before publishing, use the 'Recalculate Scores' button to:")
    bullets = [
        "Refresh all Stableford calculations based on current handicaps",
        "Update rankings to reflect any score corrections",
        "Verify that all calculations are correct before finalising",
    ]
    for b in bullets:
        p = doc.add_paragraph(b, style='List Bullet')
        for run in p.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(11)

    add_callout_box(doc, "Important", "Always click 'Recalculate Scores' before finalising, "
                    "especially if you've made any changes to scores or handicaps after "
                    "initial entry.", box_type="important")

    add_heading_styled(doc, "6.2 Finalise & Publish", level=2)
    styled_paragraph(doc,
        "When you click 'Finalise & Publish', the system performs several actions automatically:")
    styled_paragraph(doc, "")

    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(table)

    actions = [
        ("Action", "What Happens"),
        ("Generate Prizes", "Determines 1st, 2nd, 3rd place winners, plus Front 9 and Back 9 winners"),
        ("Apply Deductions", "Applies each player's deductions to their Stableford score to determine final rankings for prizes"),
        ("Auto-Write Deductions", "Automatically calculates and writes new deductions for the next event based on this event's results"),
        ("Update GOTY", "Updates the Golfer of the Year standings with this event's Stableford scores"),
    ]
    for i, (action, what) in enumerate(actions):
        row = table.rows[i]
        p0 = row.cells[0].paragraphs[0]
        r0 = p0.add_run(action)
        r0.font.name = FONT_NAME
        r0.font.size = Pt(10)
        r0.font.bold = True

        p1 = row.cells[1].paragraphs[0]
        r1 = p1.add_run(what)
        r1.font.name = FONT_NAME
        r1.font.size = Pt(10)

        if i == 0:
            set_cell_shading(row.cells[0], "1B5E3A")
            set_cell_shading(row.cells[1], "1B5E3A")
            r0.font.color.rgb = WHITE
            r1.font.color.rgb = WHITE
        elif i % 2 == 0:
            set_cell_shading(row.cells[0], "E8F5E9")
            set_cell_shading(row.cells[1], "E8F5E9")

    styled_paragraph(doc, "")

    add_heading_styled(doc, "6.3 Revert to In Progress", level=2)
    styled_paragraph(doc,
        "If you need to make changes after publishing, you can use 'Revert to In Progress'. "
        "This will:")
    bullets = [
        "Change the event status back to 'In Progress'",
        "Remove the generated prizes",
        "Reverse the deductions that were auto-written for the next event",
        "Remove this event's contribution to the GOTY standings",
        "Allow you to edit scores, add/remove players, and re-publish",
    ]
    for b in bullets:
        p = doc.add_paragraph(b, style='List Bullet')
        for run in p.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(11)

    add_callout_box(doc, "Note", "Use 'Revert to In Progress' carefully. It's designed for "
                    "corrections — for example, if you discover a scoring error after publishing. "
                    "Always re-publish after making your changes.", box_type="note")

    doc.add_page_break()

    # ── 7. Deductions Management ──
    add_heading_styled(doc, "7. Deductions Management", level=1)
    styled_paragraph(doc,
        "The deductions system is a society-specific handicapping mechanism that ensures "
        "consistent winners don't dominate every event. Understanding how it works is "
        "crucial for accurate results.")

    add_heading_styled(doc, "7.1 How Deductions Work", level=2)
    styled_paragraph(doc,
        "Each player carries a deductions total that accumulates throughout the season. "
        "This total is composed of:")
    bullets = [
        "Year starting deductions — carried over or reset at the start of each season",
        "Per-outing adjustments — added or subtracted based on results",
    ]
    for b in bullets:
        p = doc.add_paragraph(b, style='List Bullet')
        for run in p.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(11)

    add_heading_styled(doc, "7.2 Deduction Values", level=2)

    table = doc.add_table(rows=6, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(table)

    values = [
        ("Result", "Deduction Change"),
        ("1st Place", "-3 (three shots deducted)"),
        ("2nd Place", "-2 (two shots deducted)"),
        ("3rd Place", "-2 (two shots deducted)"),
        ("Front 9 / Back 9 Winner", "-1 (one shot deducted)"),
        ("Played but didn't win", "+1 (one shot earned back)"),
    ]
    for i, (result, change) in enumerate(values):
        row = table.rows[i]
        p0 = row.cells[0].paragraphs[0]
        r0 = p0.add_run(result)
        r0.font.name = FONT_NAME
        r0.font.size = Pt(10)
        r0.font.bold = (i == 0)

        p1 = row.cells[1].paragraphs[0]
        r1 = p1.add_run(change)
        r1.font.name = FONT_NAME
        r1.font.size = Pt(10)
        r1.font.bold = (i == 0)

        if i == 0:
            set_cell_shading(row.cells[0], "1B5E3A")
            set_cell_shading(row.cells[1], "1B5E3A")
            r0.font.color.rgb = WHITE
            r1.font.color.rgb = WHITE
        elif i % 2 == 0:
            set_cell_shading(row.cells[0], "E8F5E9")
            set_cell_shading(row.cells[1], "E8F5E9")

    styled_paragraph(doc, "")

    add_heading_styled(doc, "7.3 Where Deductions Apply", level=2)

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(table)

    applies = [
        ("Competition", "Deductions Applied?"),
        ("Overall Prize (1st/2nd/3rd)", "✅ Yes — deductions affect the final ranking"),
        ("Front 9 / Back 9 Prizes", "❌ No — these are based on raw Stableford points"),
        ("GOTY (Golfer of the Year)", "❌ No — GOTY uses pure Stableford scores"),
    ]
    for i, (comp, applied) in enumerate(applies):
        row = table.rows[i]
        p0 = row.cells[0].paragraphs[0]
        r0 = p0.add_run(comp)
        r0.font.name = FONT_NAME
        r0.font.size = Pt(10)
        r0.font.bold = (i == 0)

        p1 = row.cells[1].paragraphs[0]
        r1 = p1.add_run(applied)
        r1.font.name = FONT_NAME
        r1.font.size = Pt(10)
        r1.font.bold = (i == 0)

        if i == 0:
            set_cell_shading(row.cells[0], "1B5E3A")
            set_cell_shading(row.cells[1], "1B5E3A")
            r0.font.color.rgb = WHITE
            r1.font.color.rgb = WHITE
        elif i % 2 == 0:
            set_cell_shading(row.cells[0], "E8F5E9")
            set_cell_shading(row.cells[1], "E8F5E9")

    styled_paragraph(doc, "")

    add_heading_styled(doc, "7.4 Viewing and Editing the Deductions Sheet", level=2)
    styled_paragraph(doc, "To view or edit deductions:")
    steps = [
        "Open the event from the admin dashboard",
        "Go to the Deductions tab",
        "You'll see each player's deduction total",
        "To edit: click on a player's deduction value and enter the corrected amount",
        "Changes are saved automatically",
    ]
    for i, s in enumerate(steps, 1):
        p = doc.add_paragraph(f"{i}. {s}")
        for run in p.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(11)

    add_callout_box(doc, "Important", "Manual edits to deductions should be rare. The system "
                    "automatically manages deductions when you Finalise & Publish an event. "
                    "Only edit manually if correcting an error.", box_type="important")

    doc.add_page_break()

    # ── 8. Member Management ──
    add_heading_styled(doc, "8. Member Management", level=1)
    styled_paragraph(doc,
        "Member management is handled through the admin handicaps page. This is where "
        "you keep handicap indexes up to date for all society members.")

    add_heading_styled(doc, "8.1 Accessing Member Management", level=2)
    styled_paragraph(doc, "To manage members:")
    steps = [
        "Go to fairwayconnect.fly.dev/admin/handicaps",
        "You'll see a list of all registered members with their current handicap indexes",
    ]
    for i, s in enumerate(steps, 1):
        p = doc.add_paragraph(f"{i}. {s}")
        for run in p.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(11)

    add_heading_styled(doc, "8.2 Updating Handicap Indexes", level=2)
    styled_paragraph(doc, "To update a player's handicap index:")
    steps = [
        "Find the player in the list",
        "Click on their handicap index value",
        "Enter the new handicap index",
        "Save the change",
    ]
    for i, s in enumerate(steps, 1):
        p = doc.add_paragraph(f"{i}. {s}")
        for run in p.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(11)

    add_callout_box(doc, "Note", "When you update a player's handicap index, the system "
                    "automatically recalculates their playing handicap for any upcoming events "
                    "they're registered for. This ensures scores are always based on the "
                    "most current handicap.", box_type="note")

    add_callout_box(doc, "Important", "Check handicap indexes regularly against the official "
                    "Golf Ireland handicap system to ensure accuracy. Players' indexes can "
                    "change after any qualifying round.", box_type="important")

    doc.add_page_break()

    # ── 9. Printing ──
    add_heading_styled(doc, "9. Printing", level=1)
    styled_paragraph(doc,
        "FairwayConnect provides print-ready formats for key documents you may need "
        "on the day of an outing.")

    add_heading_styled(doc, "9.1 Print Scorecards", level=2)
    styled_paragraph(doc,
        "To print scorecards for an event:")
    steps = [
        "Open the event from the admin dashboard",
        "Click the 'Print Scorecards' button",
        "The system generates formatted scorecards for all confirmed players",
        "Each scorecard includes: player name, playing handicap, hole details, and stroke index",
        "Use your browser's print function (Ctrl+P / Cmd+P) to print",
    ]
    for i, s in enumerate(steps, 1):
        p = doc.add_paragraph(f"{i}. {s}")
        for run in p.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(11)

    add_heading_styled(doc, "9.2 Deductions PDF", level=2)
    styled_paragraph(doc,
        "To generate a printable deductions sheet:")
    steps = [
        "Go to the Deductions section",
        "Click the 'Print' or 'Export PDF' option",
        "The deductions sheet will be formatted for A4 printing",
        "This is useful for displaying at the society event or distributing to members",
    ]
    for i, s in enumerate(steps, 1):
        p = doc.add_paragraph(f"{i}. {s}")
        for run in p.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(11)

    add_callout_box(doc, "Tip", "Print scorecards and the deductions sheet before each outing "
                    "so players can see their current deductions on the day.",
                    box_type="tip")

    # ── End ──
    styled_paragraph(doc, "")
    doc.add_page_break()
    add_heading_styled(doc, "Quick Reference", level=1)
    styled_paragraph(doc,
        "Here's a quick reference for the most common admin tasks:")

    table = doc.add_table(rows=9, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(table)

    qr = [
        ("Task", "Where to Go"),
        ("Create event", "Admin Dashboard → Create Event"),
        ("Add players", "Event → Players tab → Add Player"),
        ("Enter scores", "Event → Scorecards tab → Select player"),
        ("Add side comps", "Event → Side Competitions"),
        ("Publish results", "Event → Recalculate Scores → Finalise & Publish"),
        ("Revert results", "Event → Revert to In Progress"),
        ("Update handicaps", "/admin/handicaps"),
    ]
    for i, (task, where) in enumerate(qr):
        row = table.rows[i]
        p0 = row.cells[0].paragraphs[0]
        r0 = p0.add_run(task)
        r0.font.name = FONT_NAME
        r0.font.size = Pt(10)
        r0.font.bold = (i == 0)

        p1 = row.cells[1].paragraphs[0]
        r1 = p1.add_run(where)
        r1.font.name = FONT_NAME
        r1.font.size = Pt(10)
        r1.font.bold = (i == 0)

        if i == 0:
            set_cell_shading(row.cells[0], "1B5E3A")
            set_cell_shading(row.cells[1], "1B5E3A")
            r0.font.color.rgb = WHITE
            r1.font.color.rgb = WHITE
        elif i % 2 == 0:
            set_cell_shading(row.cells[0], "E8F5E9")
            set_cell_shading(row.cells[1], "E8F5E9")

    # Save
    path = os.path.join(OUTPUT_DIR, 'FairwayConnect-Admin-Guide-v1.0.docx')
    doc.save(path)
    print(f"✅ Admin Guide saved to: {path}")
    return path


# ═══════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════

if __name__ == '__main__':
    print("Generating FairwayConnect User Manuals...")
    print()
    create_member_guide()
    create_admin_guide()
    print()
    print("✅ Both documents generated successfully!")
