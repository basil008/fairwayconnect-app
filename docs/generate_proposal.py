#!/usr/bin/env python3
"""Generate FairwayConnect Deployment & Landing Page Proposal .docx"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

GREEN = RGBColor(0x1B, 0x5E, 0x3A)
DARK_GREEN = RGBColor(0x14, 0x47, 0x2B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
GREY = RGBColor(0x33, 0x33, 0x33)
LIGHT_GREY = RGBColor(0x66, 0x66, 0x66)

doc = Document()

# -- Set default font to Calibri --
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = GREY

# -- Set margins --
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# -- Configure heading styles --
for i in range(1, 4):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = GREEN
    hs.font.bold = True
    if i == 1:
        hs.font.size = Pt(22)
    elif i == 2:
        hs.font.size = Pt(16)
    else:
        hs.font.size = Pt(13)
    hs.paragraph_format.space_before = Pt(18 if i == 1 else 12)
    hs.paragraph_format.space_after = Pt(6)

# Helper functions
def add_green_bar():
    """Add a green horizontal rule"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    # Use a bordered paragraph
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="12" w:space="1" w:color="1B5E3A"/></w:pBdr>')
    pPr.append(pBdr)

def add_bullet(text, bold_prefix=None, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.5 + level * 0.25)
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.color.rgb = GREY
        run2 = p.add_run(text)
        run2.font.name = 'Calibri'
        run2.font.size = Pt(11)
        run2.font.color.rgb = GREY
    else:
        p.clear()
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.color.rgb = GREY
    return p

def add_body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.color.rgb = GREY
    return p

def add_bold_body(bold_text, normal_text=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(bold_text)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = GREY
    if normal_text:
        run2 = p.add_run(normal_text)
        run2.font.name = 'Calibri'
        run2.font.size = Pt(11)
        run2.font.color.rgb = GREY
    return p

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_footer(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("FairwayConnect — Deployment Proposal")
    run.font.name = 'Calibri'
    run.font.size = Pt(9)
    run.font.color.rgb = LIGHT_GREY

# ============================================================
# TITLE PAGE
# ============================================================

# Add some spacing before title
for _ in range(6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)

# Logo text
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("⛳ FairwayConnect")
run.font.name = 'Calibri'
run.font.size = Pt(36)
run.font.color.rgb = GREEN
run.bold = True

add_green_bar()

# Subtitle
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(18)
run = p.add_run("Deployment & Landing Page Proposal")
run.font.name = 'Calibri'
run.font.size = Pt(22)
run.font.color.rgb = DARK_GREEN

# Society name
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(24)
run = p.add_run("Aer Lingus Golf Society")
run.font.name = 'Calibri'
run.font.size = Pt(18)
run.font.color.rgb = GREY

# Date
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(36)
run = p.add_run("14 April 2026")
run.font.name = 'Calibri'
run.font.size = Pt(14)
run.font.color.rgb = LIGHT_GREY

# Prepared for
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(12)
run = p.add_run("Prepared for: Basil Cooney")
run.font.name = 'Calibri'
run.font.size = Pt(14)
run.font.color.rgb = LIGHT_GREY

# Add footer to first section
add_footer(doc.sections[0])

# ============================================================
# PAGE BREAK — CONTENT STARTS
# ============================================================
doc.add_page_break()

# ============================================================
# 1. DEPLOYMENT STRATEGY
# ============================================================
doc.add_heading('1. Deployment Strategy — Getting the App to All Members', level=1)
add_green_bar()

doc.add_heading('1.1 Current Setup', level=2)
add_bullet("Web app live at fairwayconnect.fly.dev")
add_bullet("No app store needed — runs in any modern browser")
add_bullet("Works on iPhone, Android, tablet, and desktop")
add_bullet("Progressive Web App (PWA) — can be installed to the home screen for a native app feel")

doc.add_heading('1.2 Deployment Options', level=2)

# Option A
doc.add_heading('Option A: WhatsApp Group Message (Fastest)', level=3)
add_body("Send a single message to the ALGS WhatsApp group containing:")
add_bullet("Direct link to the app")
add_bullet('"Add to Home Screen" instructions for iPhone and Android')
add_bullet("Brief description of available features")
add_body("Members click the link, bookmark it, and they're in. Total time: 5 minutes.")

# Option B
doc.add_heading('Option B: QR Code Poster/Card', level=3)
add_body("Generate a QR code pointing to fairwayconnect.fly.dev and print it on:")
add_bullet("Society competition cards")
add_bullet("Small business card-sized handouts")
add_bullet("A poster at the next outing registration desk")
add_body("Members scan the QR code with their phone camera and the app opens instantly. The card can also include 'Add to Home Screen' instructions.")

# Option C
doc.add_heading('Option C: Email to All Members', level=3)
add_bullet("Send an email with the link and step-by-step instructions")
add_bullet("Good for members who aren't on WhatsApp")
add_bullet("Include screenshots showing what the app looks like")

# Recommendation
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
p.paragraph_format.space_after = Pt(6)
run = p.add_run("✅ Recommended: Option A + B combined")
run.bold = True
run.font.name = 'Calibri'
run.font.size = Pt(12)
run.font.color.rgb = GREEN
p.add_run(" — Send the WhatsApp message now for immediate access, and prepare QR code cards for the next outing.").font.name = 'Calibri'

doc.add_heading('1.3 "Add to Home Screen" Instructions', level=2)

add_bold_body("iPhone (Safari):")
add_bullet("Open fairwayconnect.fly.dev in Safari")
add_bullet("Tap the Share button (square with arrow)")
add_bullet('Scroll down and tap "Add to Home Screen"')
add_bullet("Tap Add — the app icon appears on your home screen")

add_bold_body("Android (Chrome):")
add_bullet("Open fairwayconnect.fly.dev in Chrome")
add_bullet("Tap the menu (⋮ three dots, top right)")
add_bullet('Tap "Add to Home Screen"')
add_bullet("Tap Add — the app icon appears on your home screen")

add_body("Once added, the app opens full screen — just like a native app from the App Store.")

doc.add_heading('1.4 WhatsApp Message Template', level=2)
add_body("Ready-to-send message for the ALGS WhatsApp group:")

# Message box - use a table with green border
table = doc.add_table(rows=1, cols=1)
table.style = 'Table Grid'
cell = table.cell(0, 0)
set_cell_shading(cell, "F0F8F0")

msg_text = """⛳ Aer Lingus Golf Society — FairwayConnect is LIVE!

Your new golf society app is ready. View results, leaderboards, GOTY standings, and upcoming events:

👉 https://fairwayconnect.fly.dev

To add it as an app on your phone:
📱 iPhone: Open in Safari → tap Share → "Add to Home Screen"
📱 Android: Open in Chrome → tap ⋮ → "Add to Home Screen"

No download needed — it's a web app that works in your browser!

Features:
🏆 Live Leaderboard during outings
📅 Season Calendar with all events
🏆 Golfer of the Year standings
⛳ Score entry from your phone
📊 Results and prizes

See you on the fairway! ⛳"""

p = cell.paragraphs[0]
p.clear()
run = p.add_run(msg_text)
run.font.name = 'Calibri'
run.font.size = Pt(10)
run.font.color.rgb = DARK_GREEN

# ============================================================
# 2. LANDING PAGE DESIGN
# ============================================================
doc.add_page_break()
doc.add_heading('2. Landing Page Design', level=1)
add_green_bar()

doc.add_heading('2.1 Concept', level=2)
add_body("A branded entry page that serves as the gateway to FairwayConnect. The landing page provides:")
add_bullet("FairwayConnect branding with Aer Lingus Golf Society identity")
add_bullet("Two clear entry paths: Member and Admin")
add_bullet("Simple PIN-based access (no passwords or email sign-ups)")
add_bullet("Guest access for visitors who want to browse without logging in")

doc.add_heading('2.2 Page Layout', level=2)

add_bold_body("Header Section:")
add_bullet("Large ⛳ FairwayConnect logo/wordmark")
add_bullet('"Aer Lingus Golf Society" subtitle')
add_bullet("Green gradient background (#1B5E3A to #0D3320)")
add_bullet("Clean, modern look — welcoming and easy to read")

add_bold_body("Login Section — Two Cards:")
add_body("The login area presents two cards, side by side on desktop or stacked on mobile:")

# Member Card description
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
run = p.add_run("🏌️ Member Access Card")
run.bold = True
run.font.name = 'Calibri'
run.font.size = Pt(12)
run.font.color.rgb = GREEN
add_bullet('"Member" heading')
add_bullet('"View results, leaderboards, and events"')
add_bullet("4-digit PIN input field")
add_bullet('"Continue as Guest" button for read-only access')
add_bullet("Guest access: view results, calendar, GOTY, leaderboard")
add_bullet("Member PIN: full access including score entry and personal profile")

# Admin Card description
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
run = p.add_run("⚙️ Admin Access Card")
run.bold = True
run.font.name = 'Calibri'
run.font.size = Pt(12)
run.font.color.rgb = GREEN
add_bullet('"Admin" heading')
add_bullet('"Manage events, scores, and members"')
add_bullet("4-digit PIN input field")
add_bullet("Current admin PIN: 2026")

add_bold_body("Guest Access:")
add_bullet('"Continue without login" link displayed below the cards')
add_bullet("Takes users straight to the member home page in read-only mode")
add_bullet("Can view: calendar, results, GOTY standings, leaderboard")
add_bullet("Cannot: enter scores, RSVP to events")

doc.add_heading('2.3 PIN System', level=2)

add_bold_body("Member PINs:")
add_bullet("Each member receives a unique 4-digit PIN")
add_bullet("Admin assigns PINs via the member management page")
add_bullet("PIN identifies the member — enables score entry, RSVP, and personal profile")
add_bullet("If a member forgets their PIN, the admin can reset it")

add_bold_body("Admin PIN:")
add_bullet("Single admin PIN for all admin functions")
add_bullet("Currently set to: 2026")
add_bullet("Can be changed in the admin settings")

add_bold_body("No PIN (Guest):")
add_bullet("View-only access to all public pages")
add_bullet("Ideal for prospective members, family, or casual visitors")

doc.add_heading('2.4 URL Structure', level=2)

# URL table
table = doc.add_table(rows=5, cols=2)
table.style = 'Table Grid'

headers = ['URL', 'Purpose']
data = [
    ['fairwayconnect.fly.dev', 'Landing page (root)'],
    ['fairwayconnect.fly.dev/home', 'Member home (after login)'],
    ['fairwayconnect.fly.dev/admin', 'Admin dashboard (after login)'],
    ['fairwayconnect.fly.dev/home', 'Guest view (read-only)'],
]

# Header row
for i, h in enumerate(headers):
    cell = table.cell(0, i)
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(h)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.color.rgb = WHITE
    set_cell_shading(cell, "1B5E3A")

# Data rows
for r, row_data in enumerate(data):
    for c, val in enumerate(row_data):
        cell = table.cell(r + 1, c)
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(val)
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.color.rgb = GREY

doc.add_heading('2.5 Mobile Experience', level=2)
add_bullet("Landing page is mobile-first — most members will access from their phones")
add_bullet("Large touch targets for PIN entry (easy to tap on small screens)")
add_bullet('"Remember me" option saves the PIN in the browser for auto-login next time')
add_bullet('"Add to Home Screen" prompt displayed on first visit')
add_bullet("Responsive layout adapts seamlessly from phone to tablet to desktop")

# ============================================================
# 3. SECURITY CONSIDERATIONS
# ============================================================
doc.add_page_break()
doc.add_heading('3. Security Considerations', level=1)
add_green_bar()

add_bullet("PINs are simple (4 digits) — this is appropriate for a golf society, not a bank")
add_bullet("Admin PIN gives access to score management — should be shared only with authorised administrators")
add_bullet("No sensitive personal data stored beyond names, handicaps, and scores")
add_bullet("HTTPS encryption on all connections (provided automatically by Fly.io)")
add_bullet("PIN attempts can be rate-limited to prevent brute force")
add_bullet("Session tokens expire after a configurable period for added security")

# ============================================================
# 4. IMPLEMENTATION PLAN
# ============================================================
doc.add_heading('4. Implementation Plan', level=1)
add_green_bar()

# Implementation table
table = doc.add_table(rows=6, cols=3)
table.style = 'Table Grid'

impl_headers = ['Step', 'Task', 'Time Estimate']
impl_data = [
    ['1', 'Build landing page with PIN login', '1–2 days'],
    ['2', 'Generate member PINs (automated)', '1 hour'],
    ['3', 'Deploy to Fly.io and test', '1 day'],
    ['4', 'Send WhatsApp message to ALGS group', '5 minutes'],
    ['5', 'Print QR code cards for next outing', '1 day (print)'],
]

for i, h in enumerate(impl_headers):
    cell = table.cell(0, i)
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(h)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.color.rgb = WHITE
    set_cell_shading(cell, "1B5E3A")

for r, row_data in enumerate(impl_data):
    for c, val in enumerate(row_data):
        cell = table.cell(r + 1, c)
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(val)
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.color.rgb = GREY

add_body("")
add_bold_body("Total estimated time: ", "3–4 days from start to members using the app.")

# ============================================================
# 5. QR CODE
# ============================================================
doc.add_heading('5. QR Code', level=1)
add_green_bar()

# QR code placeholder box
table = doc.add_table(rows=1, cols=1)
table.style = 'Table Grid'
cell = table.cell(0, 0)
set_cell_shading(cell, "F5F5F5")
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(24)
p.paragraph_format.space_after = Pt(24)
run = p.add_run("[ QR Code Placeholder ]")
run.font.name = 'Calibri'
run.font.size = Pt(16)
run.font.color.rgb = LIGHT_GREY
run.bold = True

p2 = cell.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("Points to: https://fairwayconnect.fly.dev")
run2.font.name = 'Calibri'
run2.font.size = Pt(11)
run2.font.color.rgb = GREEN

add_body("")
add_body("A QR code will be generated pointing to fairwayconnect.fly.dev — can be printed on competition cards, posters, or handed out at the next outing. Members simply scan with their phone camera to open the app instantly.")

add_bold_body("Suggested uses:")
add_bullet("Competition scorecards — small QR code in the corner")
add_bullet("Business card-sized handouts for new members")
add_bullet("Poster at the registration desk on outing day")
add_bullet("Email signature for society correspondence")

# ============================================================
# SAVE
# ============================================================
output_path = "/Users/abcooney/.openclaw/workspace/fairway-connect/docs/FairwayConnect-Deployment-Landing-Page-Proposal.docx"
doc.save(output_path)
print(f"✅ Document saved to: {output_path}")
